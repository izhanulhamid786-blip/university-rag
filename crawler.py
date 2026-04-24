"""
crawler.py - Autonomous Playwright crawler for Central University of Kashmir.

Example crawler_config.json:
{
  "base_url": "https://cukashmir.ac.in",
  "start_paths": ["/#/publiczone"],
  "output_dir": "data",
  "max_pages": 5000,
  "max_pdf_size_mb": 25,
  "headless": true,
  "nav_timeout_ms": 30000,
  "render_wait_max_ms": 8000,
  "dom_stable_ms": 500,
  "max_scrolls": 12,
  "max_retries": 3,
  "retry_backoff_seconds": 3,
  "min_delay_seconds": 0.2,
  "default_delay_seconds": 0.5,
  "max_delay_seconds": 8.0,
  "max_sitemap_urls": 5000,
  "allowed_domains": [
    "cukashmir.ac.in",
    "www.cukashmir.ac.in",
    "results.cukashmir.in",
    "www.results.cukashmir.in",
    "cukapi.disgenweb.in"
  ],
  "block_patterns": [
    "/wp-admin", "/login", "/signin",
    "\\.js(\\?|$)", "\\.css(\\?|$)", "\\.ico$",
    "\\.woff", "\\.ttf", "\\.mp4$", "\\.mp3$",
    "javascript:", "mailto:", "tel:", "whatsapp:",
    "facebook\\.com", "twitter\\.com", "instagram\\.com",
    "youtube\\.com", "linkedin\\.com", "google\\.com"
  ]
}
"""

from __future__ import annotations

import io
import json
import logging
import re
import shutil
import time
import hashlib
import urllib.error
import urllib.request
import urllib.robotparser
import xml.etree.ElementTree as ET
from collections import defaultdict, deque
from pathlib import Path
from datetime import datetime
from urllib.parse import parse_qsl, urlencode, urljoin, urlparse, urlunparse

from bs4 import BeautifulSoup, Tag as BS4Tag
import pdfplumber
from tqdm import tqdm
from colorama import Fore, Style, init

init(autoreset=True)


# ---------------------------------------------------------------------------
# Optional extraction dependencies
# ---------------------------------------------------------------------------

DEFAULT_TESSERACT_PATH = Path(r"C:\Users\hp\AppData\Local\Programs\Tesseract-OCR\tesseract.exe")

try:
    import pytesseract
    from pdf2image import convert_from_bytes

    _tesseract_cmd = str(DEFAULT_TESSERACT_PATH) if DEFAULT_TESSERACT_PATH.exists() else shutil.which("tesseract")
    if _tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = _tesseract_cmd
        OCR_AVAILABLE = True
        OCR_STATUS = f"enabled ({_tesseract_cmd})"
    else:
        OCR_AVAILABLE = False
        OCR_STATUS = "disabled (tesseract executable not found)"
except ImportError:
    OCR_AVAILABLE = False
    OCR_STATUS = "disabled (missing pytesseract/pdf2image)"

try:
    import docx as python_docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl

    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False


# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

CONFIG_FILE = Path("crawler_config.json")

DEFAULT_CONFIG = {
    "base_url": "https://cukashmir.ac.in",
    "start_paths": ["/#/publiczone"],
    "output_dir": "data",
    "log_file": "crawler.log",
    "max_pages": 5000,
    "max_pdf_size_mb": 25,
    "headless": True,
    "nav_timeout_ms": 30_000,
    "render_wait_max_ms": 8_000,
    "dom_stable_ms": 500,
    "max_scrolls": 12,
    "max_retries": 3,
    "retry_backoff_seconds": 3,
    "min_delay_seconds": 0.2,
    "default_delay_seconds": 0.5,
    "max_delay_seconds": 8.0,
    "max_sitemap_urls": 5000,
    "user_agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/122.0.0.0 Safari/537.36"
    ),
    "allowed_domains": [
        "cukashmir.ac.in",
        "www.cukashmir.ac.in",
        "results.cukashmir.in",
        "www.results.cukashmir.in",
        "cukapi.disgenweb.in",
    ],
    "block_patterns": [
        r"/wp-admin",
        r"/login",
        r"/signin",
        r"\.js(\?|$)",
        r"\.css(\?|$)",
        r"\.ico$",
        r"\.woff",
        r"\.ttf",
        r"\.mp4$",
        r"\.mp3$",
        r"javascript:",
        r"mailto:",
        r"tel:",
        r"whatsapp:",
        r"facebook\.com",
        r"twitter\.com",
        r"instagram\.com",
        r"youtube\.com",
        r"linkedin\.com",
        r"google\.com",
    ],
    "button_trigger_words": ["all", "more", "view", "show", "load", "browse", "explore", "read"],
    "category_synonyms": {
        "admissions": [
            "admission",
            "admissions",
            "apply",
            "application",
            "entrance",
            "cuet",
            "merit list",
            "selected candidates",
            "selection list",
            "counselling",
            "counseling",
            "prospectus",
        ],
        "fees": [
            "fee",
            "fees",
            "fee structure",
            "challan",
            "payment",
            "semester fee",
            "tuition",
            "hostel fee",
            "refund",
        ],
        "results": [
            "result",
            "results",
            "marksheet",
            "mark sheet",
            "grade card",
            "score card",
            "scorecard",
            "transcript",
            "revaluation",
            "evaluation",
            "declared",
        ],
        "examinations": [
            "exam",
            "examination",
            "date sheet",
            "date-sheet",
            "datesheet",
            "schedule",
            "timetable",
            "time table",
            "admit card",
            "hall ticket",
        ],
        "academics": [
            "academic",
            "academics",
            "syllabus",
            "course",
            "courses",
            "programme",
            "programmes",
            "program",
            "curriculum",
            "ordinance",
            "regulation",
        ],
        "faculty": ["faculty", "teacher", "professor", "assistant professor", "associate professor"],
        "departments": ["department", "departments", "school", "schools", "centre", "center"],
        "scholarships": ["scholarship", "fellowship", "stipend", "financial aid"],
        "notices": ["notice", "notification", "circular", "announcement", "office order", "order"],
        "recruitment": [
            "job",
            "jobs",
            "career",
            "careers",
            "recruitment",
            "vacancy",
            "vacancies",
            "employment",
            "employ",
            "teaching position",
            "non-teaching",
            "walk-in",
        ],
        "tenders": ["tender", "tenders", "quotation", "bid", "bids", "eoi", "procurement"],
        "news": ["news", "press release", "press releases", "media"],
        "events": ["event", "events", "seminar", "conference", "workshop", "webinar"],
        "research": ["research", "project", "projects", "publication", "publications", "phd", "doctoral"],
        "accreditation": ["naac", "nirf", "iqac", "diqa", "aqar", "accreditation", "ranking"],
        "library": ["library", "opac", "e-resource", "eresource", "e-resources", "journal", "journals"],
        "hostels": ["hostel", "hostels", "accommodation"],
        "placements": ["placement", "placements", "internship", "training"],
        "contact": ["contact", "email", "phone", "telephone", "address", "directory"],
        "downloads": ["download", "downloads", "form", "forms", "brochure", "bulletin"],
        "rti": ["rti", "right to information"],
        "about": ["about", "profile", "vision", "mission", "statute", "act"],
        "campuses": ["campus", "campuses"],
    },
}


def _deep_merge(base: dict, override: dict) -> dict:
    merged = dict(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = _deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged


def load_config(path: Path = CONFIG_FILE) -> dict:
    if not path.exists():
        path.write_text(json.dumps(DEFAULT_CONFIG, indent=2), encoding="utf-8")
    try:
        user_config = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise RuntimeError(f"Invalid {path}: {exc}") from exc
    return _deep_merge(DEFAULT_CONFIG, user_config)


CONFIG = load_config()

BASE_URL = CONFIG["base_url"].rstrip("/")
START_URLS = [normalised if (normalised := urljoin(BASE_URL + "/", p.lstrip("/"))) else BASE_URL for p in CONFIG["start_paths"]]
OUTPUT_DIR = Path(CONFIG["output_dir"])
PDF_DIR = OUTPUT_DIR / "pdfs"
JSON_DIR = OUTPUT_DIR / "structured"
LOG_FILE = CONFIG.get("log_file", "crawler.log")

MAX_PAGES = int(CONFIG["max_pages"])
MAX_PDF_SIZE = int(CONFIG["max_pdf_size_mb"]) * 1024 * 1024
HEADLESS = bool(CONFIG["headless"])
NAV_TIMEOUT_MS = int(CONFIG["nav_timeout_ms"])
RENDER_WAIT_MAX_MS = int(CONFIG["render_wait_max_ms"])
DOM_STABLE_MS = int(CONFIG["dom_stable_ms"])
MAX_SCROLLS = int(CONFIG["max_scrolls"])
MAX_RETRIES = int(CONFIG["max_retries"])
RETRY_BACKOFF_SECONDS = float(CONFIG["retry_backoff_seconds"])
ALLOWED_DOMAINS = set(CONFIG["allowed_domains"])
BLOCK_PATTERNS = list(CONFIG["block_patterns"])
BUTTON_TRIGGER_WORDS = tuple(w.lower() for w in CONFIG.get("button_trigger_words", []))
CATEGORY_SYNONYMS = CONFIG["category_synonyms"]
HIGH_VALUE_TERMS = sorted({term.lower() for terms in CATEGORY_SYNONYMS.values() for term in terms}, key=len, reverse=True)
BASE_PARTS = urlparse(BASE_URL)
BASE_ORIGIN = f"{BASE_PARTS.scheme}://{BASE_PARTS.netloc}"


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
log = logging.getLogger("cuk")


# ---------------------------------------------------------------------------
# URL, text, category, and contact helpers
# ---------------------------------------------------------------------------

EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I)
PHONE_RE = re.compile(
    r"(?<!\d)(?:\+?91[\s-]?)?(?:0[\s-]?)?(?:[6-9]\d{9}|(?:19|18|17|16|15|14|13|12|11|10|01|02|03|04|05|06|07|08|09|0194|0195)[\s-]?\d{5,8})(?!\d)"
)


def url_hash(url: str) -> str:
    return hashlib.md5(url.encode("utf-8")).hexdigest()[:12]


def normalise(url: str) -> str:
    url = (url or "").strip()
    if not url:
        return ""
    p = urlparse(url)
    if p.scheme and p.scheme not in {"http", "https"}:
        return url
    path = p.path.rstrip("/") or "/"
    return urlunparse((p.scheme, p.netloc.lower(), path, "", p.query, p.fragment))


def route_url(route: str, current_url: str | None = None) -> str:
    route = (route or "").strip().strip("'\"")
    if not route:
        return ""
    if route.startswith(("http://", "https://")):
        return normalise(route)
    if route.startswith("#/"):
        return normalise(BASE_ORIGIN + "/" + route)
    if route.startswith("/#/"):
        return normalise(BASE_ORIGIN + route)
    if route.startswith("/"):
        if "#/" in "".join(START_URLS):
            return normalise(BASE_ORIGIN + "/#" + route)
        return normalise(urljoin(BASE_URL + "/", route.lstrip("/")))
    return normalise(urljoin(current_url or BASE_URL + "/", route))


def is_allowed(url: str) -> bool:
    p = urlparse(url)
    if not p.netloc:
        return True
    return p.netloc in ALLOWED_DOMAINS or any(p.netloc.endswith("." + domain) for domain in ALLOWED_DOMAINS)


def is_blocked(url: str) -> bool:
    return any(re.search(pat, url, re.I) for pat in BLOCK_PATTERNS)


def is_binary(url: str) -> bool:
    return url.lower().split("?")[0].rsplit(".", 1)[-1] in {"pdf", "docx", "doc", "xlsx", "xls", "txt"}


def safe_filename(url: str, ext: str) -> str:
    p = urlparse(url)
    slug = re.sub(r"[^\w\-]", "_", p.path.strip("/"))[:70]
    frag = re.sub(r"[^\w\-]", "_", p.fragment)[:30]
    query = re.sub(r"[^\w\-]", "_", p.query)[:20]
    key = "_".join(part for part in [slug, frag, query] if part)
    return f"{key or 'index'}_{url_hash(url)}{ext}"


def clean_text(text: str) -> str:
    text = re.sub(r"\r\n|\r", "\n", text or "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"Original text\s*Rate this translation", "", text, flags=re.I)
    text = re.sub(r"Your feedback will be used to help improve Google Translate", "", text, flags=re.I)

    lines = []
    seen = set()
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        lines.append(line)
    return "\n".join(lines).strip()


def _term_hits(haystack: str, terms: list[str]) -> int:
    hits = 0
    for term in terms:
        pattern = r"\b" + re.escape(term.lower()).replace(r"\ ", r"\s+") + r"\b"
        if re.search(pattern, haystack):
            hits += 2 if " " in term else 1
    return hits


def categorize(url: str, title: str = "", text: str = "") -> str:
    haystack = clean_text(f"{url} {title} {text[:1000]}").lower()
    scores = {category: _term_hits(haystack, terms) for category, terms in CATEGORY_SYNONYMS.items()}
    best_category, best_score = max(scores.items(), key=lambda item: item[1])
    return best_category if best_score > 0 else "general"


def quality_score(url: str, text: str, title: str = "") -> int:
    haystack = clean_text(f"{url} {title} {text[:1500]}").lower()
    score = 0
    for category, terms in CATEGORY_SYNONYMS.items():
        hits = _term_hits(haystack, terms)
        if hits:
            score += hits
            if category in {"admissions", "fees", "results", "examinations", "notices", "recruitment"}:
                score += 2
    return score


def priority_of(url: str) -> int:
    category = categorize(url)
    if is_binary(url):
        return 0
    if category in {"admissions", "fees", "results", "examinations", "notices", "recruitment", "tenders"}:
        return 0
    if category != "general":
        return 1
    return 2


def extract_contacts(text: str) -> list[dict]:
    contacts = []
    seen = set()
    for email in EMAIL_RE.findall(text or ""):
        value = email.strip().rstrip(".,;:")
        key = ("email", value.lower())
        if key not in seen:
            seen.add(key)
            contacts.append({"type": "email", "value": value})
    for match in PHONE_RE.findall(text or ""):
        value = re.sub(r"\s+", " ", match).strip().rstrip(".,;:")
        digits = re.sub(r"\D", "", value)
        if len(digits) < 10 or len(digits) > 13:
            continue
        key = ("phone", digits)
        if key not in seen:
            seen.add(key)
            contacts.append({"type": "phone", "value": value})
    return contacts


def empty_legacy_fields() -> dict:
    return {
        "ocr": False,
        "has_table": False,
        "table_count": 0,
        "table_row_count": 0,
        "outlinks": [],
        "document_links": [],
        "notices": [],
        "quality_score": 0,
    }


def finalize_record(record: dict, retry_count: int = 0, render_time_ms: int | None = None) -> dict:
    merged = empty_legacy_fields()
    merged.update(record)
    merged.setdefault("type", "html")
    merged.setdefault("url", "")
    merged.setdefault("title", merged.get("url", ""))
    merged.setdefault("category", categorize(merged.get("url", ""), merged.get("title", ""), merged.get("text", "")))
    merged.setdefault("text", "")
    merged.setdefault("scraped_at", datetime.utcnow().isoformat())
    merged["contacts"] = merged.get("contacts") or extract_contacts(merged.get("text", ""))
    merged["retry_count"] = retry_count
    if render_time_ms is not None:
        merged["render_time_ms"] = render_time_ms
    merged["has_table"] = bool(merged.get("has_table") or merged.get("table_count", 0) > 0 or merged.get("tables"))
    merged["table_count"] = int(merged.get("table_count") or len(merged.get("tables") or []))
    if not merged.get("table_row_count"):
        merged["table_row_count"] = sum(len(tbl) for tbl in merged.get("tables") or [] if isinstance(tbl, list))
    merged["quality_score"] = int(merged.get("quality_score") or quality_score(merged["url"], merged["text"], merged["title"]))
    return merged


# ---------------------------------------------------------------------------
# PDF / DOCX / XLSX extraction
# ---------------------------------------------------------------------------

def extract_pdf(content: bytes, url: str) -> dict | None:
    rec = _pdfplumber(content, url)
    if rec is None or len(rec.get("text", "")) < 80:
        if OCR_AVAILABLE:
            log.info(f"  -> OCR fallback: {url}")
            rec = _ocr_pdf(content, url)
    return rec


def _pdfplumber(content: bytes, url: str) -> dict | None:
    try:
        pages = []
        table_count = 0
        table_row_count = 0
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            meta = pdf.metadata or {}
            for page in pdf.pages:
                txt = page.extract_text() or ""
                rows = []
                for tbl in page.extract_tables() or []:
                    table_count += 1
                    for row in tbl:
                        r = " | ".join(str(c or "").strip() for c in row)
                        if r.strip("| "):
                            rows.append(r)
                            table_row_count += 1
                combined = clean_text(txt + ("\n" + "\n".join(rows) if rows else ""))
                pages.append(combined)
        text = "\n\n--- PAGE BREAK ---\n\n".join(pages)
        title = (meta.get("Title") or "").strip() or Path(urlparse(url).path).stem.replace("-", " ").replace("_", " ").title()
        return finalize_record(
            {
                "type": "pdf",
                "url": url,
                "title": title,
                "author": (meta.get("Author") or "").strip(),
                "pages": len(pages),
                "category": categorize(url, title, text),
                "text": text,
                "ocr": False,
                "has_table": table_count > 0,
                "table_count": table_count,
                "table_row_count": table_row_count,
                "scraped_at": datetime.utcnow().isoformat(),
            }
        )
    except Exception as exc:
        log.warning(f"pdfplumber failed [{url}]: {exc}")
        return None


def _ocr_pdf(content: bytes, url: str) -> dict | None:
    try:
        images = convert_from_bytes(content, dpi=200)
        pages = []
        for img in images:
            try:
                txt = pytesseract.image_to_string(img, lang="eng+hin")
            except Exception:
                txt = pytesseract.image_to_string(img, lang="eng")
            pages.append(clean_text(txt))
        text = "\n\n--- PAGE BREAK ---\n\n".join(pages)
        if len(text.strip()) < 50:
            return None
        title = Path(urlparse(url).path).stem.replace("-", " ").replace("_", " ").title()
        return finalize_record(
            {
                "type": "pdf",
                "url": url,
                "title": title,
                "author": "",
                "pages": len(pages),
                "category": categorize(url, title, text),
                "text": text,
                "ocr": True,
                "scraped_at": datetime.utcnow().isoformat(),
            }
        )
    except Exception as exc:
        log.warning(f"OCR failed [{url}]: {exc}")
        return None


def extract_docx(content: bytes, url: str) -> dict | None:
    if not DOCX_AVAILABLE:
        return None
    try:
        doc = python_docx.Document(io.BytesIO(content))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_count = 0
        table_row_count = 0
        for tbl in doc.tables:
            table_count += 1
            for row in tbl.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells)
                if row_text.strip("| "):
                    lines.append(row_text)
                    table_row_count += 1
        text = clean_text("\n".join(lines))
        title = Path(urlparse(url).path).stem.replace("-", " ").replace("_", " ").title()
        return finalize_record(
            {
                "type": "docx",
                "url": url,
                "title": title,
                "category": categorize(url, title, text),
                "text": text,
                "has_table": table_count > 0,
                "table_count": table_count,
                "table_row_count": table_row_count,
                "scraped_at": datetime.utcnow().isoformat(),
            }
        )
    except Exception as exc:
        log.warning(f"docx failed [{url}]: {exc}")
        return None


def extract_xlsx(content: bytes, url: str) -> dict | None:
    if not XLSX_AVAILABLE:
        return None
    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
        rows = []
        table_count = 0
        for ws in wb.worksheets:
            sheet_rows = 0
            rows.append(f"Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c or "").strip() for c in row]
                if any(cells):
                    rows.append(" | ".join(cells))
                    sheet_rows += 1
            if sheet_rows:
                table_count += 1
        text = clean_text("\n".join(rows))
        title = Path(urlparse(url).path).stem.replace("-", " ").replace("_", " ").title()
        return finalize_record(
            {
                "type": "xlsx",
                "url": url,
                "title": title,
                "category": categorize(url, title, text),
                "text": text,
                "has_table": table_count > 0,
                "table_count": table_count,
                "table_row_count": max(0, len(rows) - table_count),
                "scraped_at": datetime.utcnow().isoformat(),
            }
        )
    except Exception as exc:
        log.warning(f"xlsx failed [{url}]: {exc}")
        return None


# ---------------------------------------------------------------------------
# HTML extraction from rendered Playwright HTML
# ---------------------------------------------------------------------------

def _extract_sections(soup: BeautifulSoup) -> list[dict]:
    sections = []
    candidates = soup.find_all(["article", "section", "table", "li", "tr", "div", "p"])
    for idx, el in enumerate(candidates):
        if not isinstance(el, BS4Tag):
            continue
        text = clean_text(el.get_text(separator=" "))
        if len(text) < 40 or len(text) > 3000:
            continue
        heading = ""
        h = el.find(["h1", "h2", "h3", "h4"]) if hasattr(el, "find") else None
        if h:
            heading = h.get_text(strip=True)
        sections.append(
            {
                "id": f"section-{idx}",
                "title": heading[:120],
                "hash": hashlib.md5(text.encode("utf-8")).hexdigest(),
                "text_preview": text[:250],
            }
        )
        if len(sections) >= 200:
            break
    return sections


def extract_html(html: str, url: str, render_time_ms: int | None = None) -> dict:
    soup = BeautifulSoup(html, "lxml")
    raw_text_for_contacts = soup.get_text(separator="\n")

    for tag in soup(["script", "style", "noscript", "iframe", "nav", "footer", "header"]):
        tag.decompose()

    for tag in list(soup.find_all(True)):
        try:
            if not isinstance(tag, BS4Tag) or not hasattr(tag, "attrs"):
                continue
            cls = " ".join(tag.attrs.get("class") or [])
            tid = tag.attrs.get("id") or ""
            if re.search(r"menu|sidebar|cookie|popup|banner|social|share|breadcrumb|chatbot", cls + tid, re.I):
                tag.decompose()
        except Exception:
            continue

    title = ""
    for sel in ["h1", "h2", "title"]:
        el = soup.find(sel)
        if el:
            title = el.get_text(strip=True)
            break
    title = re.sub(r"\s*[-|]\s*Central University.*", "", title, flags=re.I).strip()
    title = title or urlparse(url).fragment or url

    meta_desc = ""
    m = soup.find("meta", attrs={"name": re.compile("description", re.I)})
    if m:
        meta_desc = (m.get("content") or "").strip()

    main = (
        soup.find("main")
        or soup.find(id=re.compile(r"^(content|main|body|wrapper)$", re.I))
        or soup.find(class_=re.compile(r"(content|main.content|entry|post.body)", re.I))
        or soup.find("article")
        or soup.body
    )
    text = clean_text((main or soup).get_text(separator="\n"))

    notices = []
    for el in soup.find_all(
        ["li", "div", "p", "tr", "span"],
        class_=re.compile(r"notice|notification|circular|news|announce|update|alert|ticker|latest|result|admission", re.I),
    ):
        t = el.get_text(strip=True)
        if 20 < len(t) < 800:
            a = el.find("a", href=True)
            notices.append({"text": t[:500], "link": normalise(urljoin(url, a["href"])) if a else None})

    tables = []
    for tbl in soup.find_all("table"):
        rows = []
        for tr in tbl.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            if any(cells):
                rows.append(cells)
        if len(rows) >= 2:
            tables.append(rows)

    outlinks = set()
    document_links = set()
    for link in harvest_links_from_soup(soup, url):
        if is_binary(link):
            document_links.add(link)
        if is_allowed(link) and not is_blocked(link):
            outlinks.add(link)

    for notice in notices:
        link = (notice or {}).get("link")
        if link and not is_blocked(link):
            if is_binary(link):
                document_links.add(normalise(link))
            elif is_allowed(link):
                outlinks.add(normalise(link))

    record = {
        "type": "html",
        "url": url,
        "title": title,
        "meta_description": meta_desc,
        "category": categorize(url, title, text),
        "text": text,
        "notices": notices[:50],
        "tables": tables[:15],
        "has_table": bool(tables),
        "table_count": len(tables),
        "table_row_count": sum(len(tbl) for tbl in tables),
        "outlinks": sorted(outlinks),
        "document_links": sorted(document_links),
        "contacts": extract_contacts(raw_text_for_contacts + "\n" + text),
        "sections": _extract_sections(soup),
        "quality_score": quality_score(url, text, title),
        "scraped_at": datetime.utcnow().isoformat(),
    }
    return finalize_record(record, render_time_ms=render_time_ms)


def harvest_links_from_soup(soup: BeautifulSoup, current_url: str) -> set[str]:
    links = set()
    attr_names = ["href", "routerlink", "ng-reflect-router-link", "data-router-link", "data-href"]
    for el in soup.find_all(True):
        for attr in attr_names:
            raw = el.get(attr)
            if not raw:
                continue
            for value in raw if isinstance(raw, list) else [raw]:
                value = str(value).strip()
                if not value or is_blocked(value):
                    continue
                abs_url = route_url(value, current_url) if "router" in attr else normalise(urljoin(current_url, value))
                if is_allowed(abs_url) and not is_blocked(abs_url):
                    links.add(abs_url)
    return links


# ---------------------------------------------------------------------------
# Rendering, rate limiting, and Playwright handling
# ---------------------------------------------------------------------------

class AdaptiveRateLimiter:
    def __init__(self):
        self.delay = float(CONFIG["default_delay_seconds"])
        self.min_delay = float(CONFIG["min_delay_seconds"])
        self.max_delay = float(CONFIG["max_delay_seconds"])
        self.crawl_delay = None
        self.recent = deque(maxlen=20)
        self.backoff_until = 0.0

    def set_crawl_delay(self, seconds: float | None):
        if seconds is not None:
            self.crawl_delay = max(self.min_delay, float(seconds))
            self.delay = max(self.delay, self.crawl_delay)

    def wait(self):
        now = time.time()
        if now < self.backoff_until:
            time.sleep(self.backoff_until - now)
        time.sleep(max(self.delay, self.crawl_delay or 0))

    def record(self, elapsed_ms: float | None = None, status: int | None = None):
        if status in {429, 503}:
            self.backoff_until = time.time() + 30
            self.delay = min(self.max_delay, max(self.delay * 2, 2.0))
            log.warning(f"Server returned {status}; backing off for 30 seconds")
            return
        if elapsed_ms is None:
            return
        self.recent.append(elapsed_ms)
        avg = sum(self.recent) / len(self.recent)
        if avg < 500:
            self.delay = max(self.min_delay, self.delay * 0.85)
        elif avg > 2000:
            self.delay = min(self.max_delay, self.delay * 1.5)


class PageHandler:
    """
    Wraps Playwright page with dynamic Angular SPA discovery helpers.
    """

    def __init__(self, playwright, rate_limiter: AdaptiveRateLimiter):
        self.rate_limiter = rate_limiter
        self.browser = playwright.chromium.launch(headless=HEADLESS)
        self.ctx = self.browser.new_context(user_agent=CONFIG["user_agent"], java_script_enabled=True)
        self.page = self.ctx.new_page()
        self.intercepted_urls: set[str] = set()
        self.request_started: dict[str, float] = {}
        self.last_render_time_ms = 0

        self.page.route(
            "**/*",
            lambda r: r.abort()
            if r.request.resource_type in ("image", "font", "media", "stylesheet")
            else r.continue_(),
        )
        self.page.on("request", self._on_request)
        self.page.on("response", self._on_response)

    def _on_request(self, request):
        url = request.url
        self.request_started[url] = time.perf_counter()
        lower_url = url.lower()
        if any(ext in lower_url for ext in (".pdf", ".docx", ".xlsx", ".xls", ".doc", ".txt")):
            if is_allowed(url) and not is_blocked(url):
                self.intercepted_urls.add(normalise(url))

    def _on_response(self, response):
        started = self.request_started.pop(response.url, None)
        elapsed = (time.perf_counter() - started) * 1000 if started else None
        self.rate_limiter.record(elapsed, response.status)

    def goto(self, url: str) -> bool:
        try:
            self.rate_limiter.wait()
            started = time.perf_counter()
            self.page.goto(url, timeout=NAV_TIMEOUT_MS, wait_until="domcontentloaded")
            self.wait_for_render()
            self.last_render_time_ms = int((time.perf_counter() - started) * 1000)
            return True
        except Exception as exc:
            log.warning(f"goto failed [{url}]: {exc}")
            return False

    def wait_for_render(self):
        try:
            self.page.wait_for_load_state("networkidle", timeout=min(RENDER_WAIT_MAX_MS, NAV_TIMEOUT_MS))
        except Exception:
            pass
        try:
            self.page.evaluate(
                """
                ({stableMs, maxMs}) => new Promise(resolve => {
                    let done = false;
                    let stableTimer = null;
                    const finish = () => {
                        if (done) return;
                        done = true;
                        observer.disconnect();
                        clearTimeout(stableTimer);
                        resolve(true);
                    };
                    const observer = new MutationObserver(() => {
                        clearTimeout(stableTimer);
                        stableTimer = setTimeout(finish, stableMs);
                    });
                    observer.observe(document.documentElement, {childList: true, subtree: true});
                    stableTimer = setTimeout(finish, stableMs);
                    setTimeout(finish, maxMs);
                })
                """,
                {"stableMs": DOM_STABLE_MS, "maxMs": RENDER_WAIT_MAX_MS},
            )
        except Exception:
            self.page.wait_for_timeout(min(DOM_STABLE_MS, 1000))

    def html(self) -> str:
        return self.page.content()

    def current_url(self) -> str:
        return normalise(self.page.url)

    def scroll_to_bottom(self) -> int:
        scrolls = 0
        stable_rounds = 0
        previous = self._page_signature()
        for _ in range(MAX_SCROLLS):
            try:
                self.page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                self.wait_for_render()
                scrolls += 1
                current = self._page_signature()
                if current == previous:
                    stable_rounds += 1
                else:
                    stable_rounds = 0
                previous = current
                if stable_rounds >= 2:
                    break
            except Exception:
                break
        return scrolls

    def _page_signature(self) -> tuple[int, int, int]:
        try:
            return self.page.evaluate(
                """
                () => [
                    document.body.scrollHeight,
                    document.querySelectorAll('a,[routerLink],[ng-reflect-router-link],button').length,
                    (document.body.innerText || '').length
                ]
                """
            )
        except Exception:
            return (0, 0, 0)

    def discover_navigation_targets(self) -> set[str]:
        links = set()
        try:
            soup = BeautifulSoup(self.html(), "lxml")
            nav_areas = soup.select("nav, header, .navbar, .nav, .menu, [role='navigation']")
            for area in nav_areas:
                links.update(harvest_links_from_soup(BeautifulSoup(str(area), "lxml"), self.current_url()))
        except Exception:
            pass
        return links

    def click_dynamic_nav(self) -> set[str]:
        found = set()
        selectors = "nav a, nav button, header a, header button, .navbar a, .navbar button, .nav a, .nav button, [role='navigation'] a, [role='navigation'] button"
        try:
            items = self.page.locator(selectors)
            count = min(items.count(), 80)
            for idx in range(count):
                item = items.nth(idx)
                try:
                    text = clean_text(item.inner_text(timeout=1200))[:120]
                    href = item.get_attribute("href", timeout=1200)
                    router = (
                        item.get_attribute("routerlink", timeout=1200)
                        or item.get_attribute("routerLink", timeout=1200)
                        or item.get_attribute("ng-reflect-router-link", timeout=1200)
                    )
                    if href or router:
                        link = route_url(router, self.current_url()) if router else normalise(urljoin(self.current_url(), href))
                        if is_allowed(link) and not is_blocked(link):
                            found.add(link)
                        continue
                    if not text:
                        continue
                    before = self.current_url()
                    item.click(timeout=2500)
                    self.wait_for_render()
                    after = self.current_url()
                    if after and after != before:
                        found.add(after)
                        self.page.goto(before, timeout=NAV_TIMEOUT_MS, wait_until="domcontentloaded")
                        self.wait_for_render()
                except Exception:
                    continue
        except Exception:
            pass
        return found

    def click_expandable_buttons(self) -> set[str]:
        found = set()
        selector = "a, button, [role='button'], span.clickable, li.clickable, .btn, .button"
        try:
            controls = self.page.locator(selector)
            count = min(controls.count(), 160)
            for idx in range(count):
                control = controls.nth(idx)
                try:
                    text_parts = [
                        control.inner_text(timeout=1000) or "",
                        control.get_attribute("aria-label", timeout=1000) or "",
                        control.get_attribute("title", timeout=1000) or "",
                    ]
                    text = clean_text(" ".join(text_parts)).lower()
                    if not text or len(text) > 90:
                        continue
                    if not any(word in text for word in BUTTON_TRIGGER_WORDS):
                        continue
                    href = control.get_attribute("href", timeout=1000)
                    router = (
                        control.get_attribute("routerlink", timeout=1000)
                        or control.get_attribute("routerLink", timeout=1000)
                        or control.get_attribute("ng-reflect-router-link", timeout=1000)
                    )
                    if href or router:
                        link = route_url(router, self.current_url()) if router else normalise(urljoin(self.current_url(), href))
                        if is_allowed(link) and not is_blocked(link):
                            found.add(link)
                        continue
                    before = self.current_url()
                    before_sig = self._page_signature()
                    control.click(timeout=2500)
                    self.wait_for_render()
                    after = self.current_url()
                    if after and after != before:
                        found.add(after)
                        self.page.goto(before, timeout=NAV_TIMEOUT_MS, wait_until="domcontentloaded")
                        self.wait_for_render()
                    elif self._page_signature() != before_sig:
                        found.update(self.harvest_all_links())
                except Exception:
                    continue
        except Exception:
            pass
        return found

    def harvest_all_links(self) -> set[str]:
        try:
            return harvest_links_from_soup(BeautifulSoup(self.html(), "lxml"), self.current_url())
        except Exception:
            return set()

    def harvest_pagination_links(self) -> set[str]:
        links = set()
        current = self.current_url()
        try:
            soup = BeautifulSoup(self.html(), "lxml")
            for a in soup.find_all("a", href=True):
                text = a.get_text(" ", strip=True).lower()
                rel = " ".join(a.get("rel") or []).lower()
                cls = " ".join(a.get("class") or []).lower()
                href = a["href"].strip()
                if rel == "next" or re.search(r"\b(next|previous|prev|older|newer)\b", text + " " + cls):
                    link = normalise(urljoin(current, href))
                    if is_allowed(link) and not is_blocked(link):
                        links.add(link)
                if re.fullmatch(r"\d{1,3}", text) and ("page" in href.lower() or "p=" in href.lower()):
                    link = normalise(urljoin(current, href))
                    if is_allowed(link) and not is_blocked(link):
                        links.add(link)
            links.update(generate_next_page_candidates(current))
        except Exception:
            pass
        return links

    def fetch_bytes(self, url: str) -> tuple[bytes | None, int | None]:
        headers = {
            "User-Agent": CONFIG["user_agent"],
            "Referer": BASE_URL,
            "Accept": "application/pdf,application/octet-stream,*/*",
        }
        started = time.perf_counter()
        try:
            self.rate_limiter.wait()
            response = self.ctx.request.get(url, headers=headers, timeout=30_000)
            elapsed = (time.perf_counter() - started) * 1000
            self.rate_limiter.record(elapsed, response.status)
            if response.ok:
                body = response.body()
                status = response.status
                response.dispose()
                return body, status
            status = response.status
            response.dispose()
            return None, status
        except Exception:
            pass

        try:
            self.rate_limiter.wait()
            started = time.perf_counter()
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req, timeout=30) as resp:
                body = resp.read()
                elapsed = (time.perf_counter() - started) * 1000
                self.rate_limiter.record(elapsed, getattr(resp, "status", None))
                return body, getattr(resp, "status", None)
        except urllib.error.HTTPError as exc:
            self.rate_limiter.record(None, exc.code)
            log.warning(f"Binary fetch failed [{url}]: HTTP {exc.code}")
            return None, exc.code
        except Exception as exc:
            log.warning(f"Binary fetch failed [{url}]: {exc}")
            return None, None

    def close(self):
        self.ctx.close()
        self.browser.close()


def generate_next_page_candidates(url: str) -> set[str]:
    candidates = set()
    p = urlparse(url)
    path_match = re.search(r"(.*/page/)(\d+)(/?)$", p.path)
    if path_match:
        next_path = f"{path_match.group(1)}{int(path_match.group(2)) + 1}{path_match.group(3)}"
        candidates.add(normalise(urlunparse((p.scheme, p.netloc, next_path, "", p.query, p.fragment))))
    query = dict(parse_qsl(p.query, keep_blank_values=True))
    for key in ("page", "p", "paged"):
        if key in query and query[key].isdigit():
            query[key] = str(int(query[key]) + 1)
            candidates.add(normalise(urlunparse((p.scheme, p.netloc, p.path, "", urlencode(query), p.fragment))))
    return candidates


# ---------------------------------------------------------------------------
# Crawler orchestration
# ---------------------------------------------------------------------------

class UniversityCrawler:
    def __init__(self, fresh_start: bool = False):
        for directory in [OUTPUT_DIR, PDF_DIR, JSON_DIR]:
            directory.mkdir(parents=True, exist_ok=True)

        self.visited: set[str] = set()
        self.permanent_failures: dict[str, dict] = {}
        self.queue: list[tuple[int, float, int, str]] = []
        self.queued: set[str] = set()
        self.sequence = 0
        self.retry_counts: defaultdict[str, int] = defaultdict(int)
        self.rate_limiter = AdaptiveRateLimiter()
        self.robot_parser: urllib.robotparser.RobotFileParser | None = None

        self.stats = {
            "html": 0,
            "pdf": 0,
            "pdf_ocr": 0,
            "pdf_failed": 0,
            "pdf_fetch_failed": 0,
            "docx": 0,
            "xlsx": 0,
            "txt": 0,
            "skipped": 0,
            "errors": 0,
            "retries": 0,
            "failed_permanent": 0,
            "contacts": 0,
            "sitemap_urls": 0,
            "pagination_urls": 0,
            "chars": 0,
        }

        self._vfile = OUTPUT_DIR / ".visited_urls.txt"
        self._failed_file = OUTPUT_DIR / ".failed_urls.json"
        self._hash_store = OUTPUT_DIR / ".doc_hashes.json"
        self._section_hash_store = OUTPUT_DIR / ".section_hashes.json"
        self._hashes: dict[str, str] = self._read_json(self._hash_store, {})
        self._section_hashes: dict[str, dict] = self._read_json(self._section_hash_store, {})
        self.permanent_failures = self._read_json(self._failed_file, {})

        if fresh_start:
            if self._vfile.exists():
                self._vfile.unlink()
            log.info("Fresh start - cleared visited URLs")

        if self._vfile.exists():
            self.visited = set(self._vfile.read_text(encoding="utf-8").splitlines())
            log.info(f"Resumed - {len(self.visited)} URLs already visited")

    @staticmethod
    def _read_json(path: Path, default):
        if not path.exists():
            return default
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return default

    def _enqueue(self, url: str, delay_seconds: float = 0.0):
        url = normalise(url)
        if not url or url in self.visited or url in self.queued or url in self.permanent_failures:
            return
        if is_blocked(url) or not is_allowed(url) or not self._robots_allowed(url):
            return
        self.sequence += 1
        self.queue.append((priority_of(url), time.time() + delay_seconds, self.sequence, url))
        self.queued.add(url)

    def _next_url(self) -> str | None:
        if not self.queue:
            return None
        now = time.time()
        ready = [item for item in self.queue if item[1] <= now]
        if not ready:
            sleep_for = max(0.0, min(item[1] for item in self.queue) - now)
            time.sleep(min(sleep_for, 5.0))
            return self._next_url()
        item = min(ready, key=lambda x: (x[0], x[1], x[2]))
        self.queue.remove(item)
        self.queued.discard(item[3])
        return item[3]

    def _mark_visited(self, url: str):
        if url in self.visited:
            return
        self.visited.add(url)
        with open(self._vfile, "a", encoding="utf-8") as f:
            f.write(url + "\n")

    def _schedule_retry(self, url: str, reason: str):
        self.retry_counts[url] += 1
        attempt = self.retry_counts[url]
        if attempt >= MAX_RETRIES:
            self.permanent_failures[url] = {
                "url": url,
                "reason": reason,
                "retry_count": attempt,
                "failed_at": datetime.utcnow().isoformat(),
            }
            self._failed_file.write_text(json.dumps(self.permanent_failures, indent=2), encoding="utf-8")
            self.stats["failed_permanent"] += 1
            self.stats["errors"] += 1
            log.error(f"Permanent failure after {attempt} attempts [{url}]: {reason}")
            return
        delay = RETRY_BACKOFF_SECONDS * (2 ** (attempt - 1))
        self.stats["retries"] += 1
        log.warning(f"Retry {attempt}/{MAX_RETRIES} in {delay:.1f}s [{url}]: {reason}")
        self._enqueue(url, delay_seconds=delay)

    def _changed(self, fname: str, content: bytes | str) -> bool:
        raw = content if isinstance(content, bytes) else content.encode("utf-8", errors="ignore")
        h = hashlib.md5(raw).hexdigest()
        if self._hashes.get(fname) == h:
            return False
        self._hashes[fname] = h
        self._hash_store.write_text(json.dumps(self._hashes, indent=2), encoding="utf-8")
        return True

    def _section_changes(self, url: str, record: dict) -> dict:
        current = {section["id"]: section["hash"] for section in record.get("sections", []) if section.get("id")}
        previous = self._section_hashes.get(url, {})
        changed = [sid for sid, digest in current.items() if previous.get(sid) != digest]
        removed = [sid for sid in previous if sid not in current]
        self._section_hashes[url] = current
        self._section_hash_store.write_text(json.dumps(self._section_hashes, indent=2), encoding="utf-8")
        return {"changed_sections": changed, "removed_sections": removed, "changed_section_count": len(changed)}

    def _save(self, record: dict):
        fname = safe_filename(record["url"], ".json")
        (JSON_DIR / fname).write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")

    def _save_html_page(self, url: str, handler: PageHandler) -> bool:
        html = handler.html()
        if not html:
            return False
        if not self._changed(safe_filename(url, ".cache"), html):
            self.stats["skipped"] += 1
            self._mark_visited(url)
            return True

        record = extract_html(html, url, handler.last_render_time_ms)
        section_delta = self._section_changes(url, record)
        record.update(section_delta)
        if len(record["text"]) < 60 and not record.get("document_links"):
            self.stats["skipped"] += 1
            self._mark_visited(url)
            return True

        self._save(record)
        self._mark_visited(url)
        self.stats["html"] += 1
        self.stats["chars"] += len(record["text"])
        self.stats["contacts"] += len(record.get("contacts", []))

        color = {
            "notices": Fore.YELLOW,
            "admissions": Fore.GREEN,
            "faculty": Fore.BLUE,
            "academics": Fore.MAGENTA,
            "fees": Fore.RED,
            "results": Fore.CYAN,
            "recruitment": Fore.LIGHTRED_EX,
        }.get(record["category"], Fore.WHITE)
        log.info(
            f"{color}[{record['category'].upper()[:12]}]{Style.RESET_ALL} "
            f"Q={record['quality_score']} | sections+{record.get('changed_section_count', 0)} | {record['title'][:55]}"
        )

        for link in record["outlinks"]:
            self._enqueue(link)
        for link in record.get("document_links", []):
            self._enqueue(link)
        return True

    def _save_binary(self, url: str, handler: PageHandler) -> bool:
        ext = url.lower().split("?")[0].rsplit(".", 1)[-1]
        content, status = handler.fetch_bytes(url)
        if not content:
            if status in {429, 503} or status is None:
                self._schedule_retry(url, f"binary fetch failed status={status}")
                if ext == "pdf":
                    self.stats["pdf_fetch_failed"] += 1
                return False
            self.stats["errors"] += 1
            self._mark_visited(url)
            return True
        if len(content) > MAX_PDF_SIZE:
            self.stats["skipped"] += 1
            self._mark_visited(url)
            return True
        if not self._changed(safe_filename(url, ".bin"), content):
            self.stats["skipped"] += 1
            self._mark_visited(url)
            return True

        record = None
        if ext == "pdf":
            record = extract_pdf(content, url)
            if record:
                (PDF_DIR / safe_filename(url, ".pdf")).write_bytes(content)
                if record.get("ocr"):
                    self.stats["pdf_ocr"] += 1
                    log.info(f"{Fore.LIGHTBLUE_EX}[PDF-OCR]{Style.RESET_ALL} {record['title'][:55]}")
                else:
                    self.stats["pdf"] += 1
                    log.info(f"{Fore.CYAN}[PDF]{Style.RESET_ALL} {record['title'][:55]} ({record.get('pages', 0)}pp)")
            else:
                self.stats["pdf_failed"] += 1
                log.warning(f"PDF extraction failed [{url}]")
        elif ext in {"docx", "doc"}:
            record = extract_docx(content, url)
            if record:
                self.stats["docx"] += 1
                log.info(f"{Fore.LIGHTGREEN_EX}[DOCX]{Style.RESET_ALL} {record['title'][:55]}")
        elif ext in {"xlsx", "xls"}:
            record = extract_xlsx(content, url)
            if record:
                self.stats["xlsx"] += 1
                log.info(f"{Fore.LIGHTGREEN_EX}[XLSX]{Style.RESET_ALL} {record['title'][:55]}")
        elif ext == "txt":
            try:
                text = clean_text(content.decode("utf-8", errors="replace"))
                if len(text) > 50:
                    title = Path(urlparse(url).path).stem
                    record = finalize_record(
                        {
                            "type": "txt",
                            "url": url,
                            "title": title,
                            "category": categorize(url, title, text),
                            "text": text,
                            "scraped_at": datetime.utcnow().isoformat(),
                        }
                    )
                    self.stats["txt"] += 1
                    log.info(f"[TXT] {title[:55]}")
            except Exception:
                pass

        if record:
            record["retry_count"] = self.retry_counts.get(url, 0)
            self._save(record)
            self.stats["chars"] += len(record.get("text", ""))
            self.stats["contacts"] += len(record.get("contacts", []))
        else:
            self.stats["skipped"] += 1
        self._mark_visited(url)
        return True

    def _load_robots(self):
        robots_url = urljoin(BASE_URL + "/", "robots.txt")
        rp = urllib.robotparser.RobotFileParser()
        rp.set_url(robots_url)
        try:
            rp.read()
            self.robot_parser = rp
            self.rate_limiter.set_crawl_delay(rp.crawl_delay("*"))
            log.info(f"robots.txt loaded; crawl-delay={rp.crawl_delay('*')}")
        except Exception as exc:
            log.warning(f"robots.txt unavailable: {exc}")
            self.robot_parser = None

    def _robots_allowed(self, url: str) -> bool:
        if not self.robot_parser:
            return True
        try:
            return self.robot_parser.can_fetch(CONFIG["user_agent"], url) or self.robot_parser.can_fetch("*", url)
        except Exception:
            return True

    def _fetch_url_text(self, url: str) -> str | None:
        try:
            req = urllib.request.Request(url, headers={"User-Agent": CONFIG["user_agent"]})
            with urllib.request.urlopen(req, timeout=20) as resp:
                return resp.read().decode("utf-8", errors="replace")
        except Exception:
            return None

    def _seed_sitemaps(self):
        seen_sitemaps = set()
        sitemap_urls = [
            urljoin(BASE_URL + "/", "sitemap.xml"),
            urljoin(BASE_URL + "/", "sitemap_index.xml"),
        ]
        total = 0
        while sitemap_urls and total < int(CONFIG["max_sitemap_urls"]):
            sitemap_url = sitemap_urls.pop(0)
            if sitemap_url in seen_sitemaps:
                continue
            seen_sitemaps.add(sitemap_url)
            xml_text = self._fetch_url_text(sitemap_url)
            if not xml_text:
                continue
            try:
                root = ET.fromstring(xml_text.encode("utf-8"))
            except Exception:
                continue
            locs = [el.text.strip() for el in root.iter() if el.tag.lower().endswith("loc") and el.text]
            for loc in locs:
                loc = normalise(loc)
                if loc.lower().endswith(".xml") and "sitemap" in loc.lower():
                    sitemap_urls.append(loc)
                else:
                    self._enqueue(loc)
                    total += 1
                    if total >= int(CONFIG["max_sitemap_urls"]):
                        break
        self.stats["sitemap_urls"] = total
        if total:
            log.info(f"Seeded {total} URLs from sitemap files")

    def _interactive_discovery(self, handler: PageHandler):
        log.info("\n--- Phase 1: Autonomous interactive discovery ---")
        for start_url in START_URLS:
            if not handler.goto(start_url):
                self._schedule_retry(start_url, "initial navigation failed")
                continue
            handler.scroll_to_bottom()
            self._save_html_page(start_url, handler)

            for link in handler.discover_navigation_targets():
                self._enqueue(link)
            for link in handler.click_dynamic_nav():
                self._enqueue(link)
            for link in handler.click_expandable_buttons():
                self._enqueue(link)
            for link in handler.harvest_all_links():
                self._enqueue(link)
            for link in handler.harvest_pagination_links():
                self.stats["pagination_urls"] += 1
                self._enqueue(link)
            for intercepted in list(handler.intercepted_urls):
                self._enqueue(intercepted)
            handler.intercepted_urls.clear()

        log.info(f"--- Discovery complete. Queue size: {len(self.queue)} ---\n")

    def _process_html_url(self, url: str, handler: PageHandler) -> bool:
        if not handler.goto(url):
            self._schedule_retry(url, "navigation failed")
            return False
        handler.scroll_to_bottom()
        if not self._save_html_page(url, handler):
            self._schedule_retry(url, "empty rendered HTML")
            return False
        for link in handler.click_expandable_buttons():
            self._enqueue(link)
        for link in handler.harvest_all_links():
            self._enqueue(link)
        for link in handler.harvest_pagination_links():
            self.stats["pagination_urls"] += 1
            self._enqueue(link)
        for intercepted in list(handler.intercepted_urls):
            self._enqueue(intercepted)
        handler.intercepted_urls.clear()
        return True

    def run(self):
        from playwright.sync_api import sync_playwright

        log.info(f"\n{'=' * 60}")
        log.info("  Autonomous University Playwright Crawler v5")
        log.info(f"  Site    : {BASE_URL}")
        log.info(f"  Starts  : {', '.join(START_URLS)}")
        log.info(f"  Output  : {OUTPUT_DIR.resolve()}")
        log.info(f"  OCR     : {OCR_STATUS}")
        log.info(f"{'=' * 60}\n")

        self._load_robots()
        self._seed_sitemaps()
        for start_url in START_URLS:
            self._enqueue(start_url)

        pbar = tqdm(total=MAX_PAGES, desc="Crawling", unit="pg")
        crawled = 0
        with sync_playwright() as pw:
            handler = PageHandler(pw, self.rate_limiter)
            try:
                self._interactive_discovery(handler)
                log.info("--- Phase 2: Processing discovered queue ---")
                while self.queue and crawled < MAX_PAGES:
                    url = self._next_url()
                    if not url or url in self.visited:
                        continue
                    if is_binary(url):
                        self._save_binary(url, handler)
                    else:
                        self._process_html_url(url, handler)

                    crawled += 1
                    pbar.update(1)
                    pbar.set_description(
                        f"HTML:{self.stats['html']} PDF:{self.stats['pdf']} "
                        f"OCR:{self.stats['pdf_ocr']} Retry:{self.stats['retries']}"
                    )
            except KeyboardInterrupt:
                log.info("\nStopped by user - progress saved.")
            finally:
                handler.close()

        pbar.close()
        self._summary()
        self._build_index()

    def _summary(self):
        print(f"\n{Fore.GREEN}{'=' * 48}")
        print("  CRAWL COMPLETE")
        print(f"{'=' * 48}{Style.RESET_ALL}")
        print(f"  HTML pages     : {self.stats['html']}")
        print(f"  PDFs (text)    : {self.stats['pdf']}")
        print(f"  PDFs (OCR)     : {self.stats['pdf_ocr']}")
        print(f"  PDFs failed    : {self.stats['pdf_failed']}")
        print(f"  PDF fetch      : {self.stats['pdf_fetch_failed']} failed")
        print(f"  DOCX           : {self.stats['docx']}")
        print(f"  XLSX           : {self.stats['xlsx']}")
        print(f"  TXT            : {self.stats['txt']}")
        print(f"  Contacts       : {self.stats['contacts']}")
        print(f"  Sitemap URLs   : {self.stats['sitemap_urls']}")
        print(f"  Pagination URLs: {self.stats['pagination_urls']}")
        print(f"  Retries        : {self.stats['retries']}")
        print(f"  Permanent fail : {self.stats['failed_permanent']}")
        print(f"  Skipped        : {self.stats['skipped']}")
        print(f"  Errors         : {self.stats['errors']}")
        print(f"  Total chars    : {self.stats['chars']:,}")
        print(f"  Output         : {JSON_DIR.resolve()}\n")

    def _build_index(self):
        index = []
        for f in JSON_DIR.glob("*.json"):
            try:
                d = json.loads(f.read_text(encoding="utf-8"))
                index.append(
                    {
                        "file": f.name,
                        "url": d.get("url"),
                        "title": d.get("title"),
                        "type": d.get("type"),
                        "category": d.get("category", "general"),
                        "quality_score": d.get("quality_score", 0),
                        "ocr": d.get("ocr", False),
                        "scraped_at": d.get("scraped_at"),
                    }
                )
            except Exception:
                pass
        index.sort(key=lambda x: x.get("quality_score", 0), reverse=True)
        idx = OUTPUT_DIR / "index.json"
        idx.write_text(json.dumps(index, indent=2), encoding="utf-8")
        log.info(f"Index saved -> {idx} ({len(index)} docs)")


CUKCrawler = UniversityCrawler


# ---------------------------------------------------------------------------
# Scheduler and entry point
# ---------------------------------------------------------------------------

def start_scheduler():
    try:
        from apscheduler.schedulers.blocking import BlockingScheduler
    except ImportError:
        print("Run: pip install apscheduler")
        return

    scheduler = BlockingScheduler()

    @scheduler.scheduled_job("cron", day_of_week="sun", hour=2)
    def job():
        log.info("=== Scheduled weekly re-crawl ===")
        UniversityCrawler(fresh_start=False).run()

    log.info("Scheduler running - re-crawls every Sunday 02:00. Ctrl+C to stop.")
    scheduler.start()


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Autonomous university crawler")
    parser.add_argument("--schedule", action="store_true", help="Run the weekly scheduler")
    parser.add_argument("--fresh", action="store_true", help="Ignore saved visited URLs and start fresh")
    args = parser.parse_args()

    if args.schedule:
        start_scheduler()
    else:
        UniversityCrawler(fresh_start=args.fresh).run()
